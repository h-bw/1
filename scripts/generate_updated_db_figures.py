from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT_DIR = Path(r"D:\DevelopSoft\lingshi\generated_figures")
FIG33 = OUT_DIR / "图3-3核心交易模块数据库表关系图-更新版.png"
FIG36 = OUT_DIR / "图3-6算法推荐模块数据库表关系图-更新版.png"
DOCX_OUT = Path(r"D:\DevelopSoft\lingshi\图3-3与图3-6-更新版.docx")

FONT_CN = r"C:\Windows\Fonts\simsun.ttc"
FONT_BOLD = r"C:\Windows\Fonts\msyhbd.ttc"

LINE = "#666666"
HEADER = "#efefef"
NOTE = "#fffbd6"


def ft(size, bold=False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_CN, size)


def multiline_size(draw, text, font, spacing=4):
    box = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing)
    return box[2] - box[0], box[3] - box[1]


def draw_center_text(draw, box, text, font, fill="black", spacing=4):
    x1, y1, x2, y2 = box
    w, h = multiline_size(draw, text, font, spacing)
    draw.multiline_text(
        ((x1 + x2 - w) / 2, (y1 + y2 - h) / 2),
        text,
        font=font,
        fill=fill,
        spacing=spacing,
        align="center",
    )


def draw_table(draw, x, y, w, title_cn, title_en, fields):
    header_h = 68
    row_h = 24
    body_h = row_h * len(fields) + 8
    h = header_h + body_h

    draw.rectangle((x, y, x + w, y + h), outline=LINE, width=1, fill="white")
    draw.rectangle((x, y, x + w, y + header_h), outline=LINE, width=1, fill=HEADER)
    draw.line((x, y + header_h, x + w, y + header_h), fill=LINE, width=1)

    draw_center_text(draw, (x, y + 6, x + w, y + 34), title_cn, ft(18, bold=True))
    draw_center_text(draw, (x, y + 30, x + w, y + 58), title_en, ft(16))

    current_y = y + header_h + 5
    for field in fields:
        draw_center_text(draw, (x + 4, current_y, x + w - 4, current_y + row_h), field, ft(14))
        current_y += row_h

    return {
        "x": x,
        "y": y,
        "w": w,
        "h": h,
        "left": (x, y + h / 2),
        "right": (x + w, y + h / 2),
        "top": (x + w / 2, y),
        "bottom": (x + w / 2, y + h),
    }


def draw_note(draw, x, y, w, h, text):
    pts = [(x, y), (x + w - 16, y), (x + w, y + 16), (x + w, y + h), (x, y + h)]
    draw.polygon(pts, outline=LINE, fill=NOTE)
    draw.line((x + w - 16, y, x + w - 16, y + 16), fill=LINE, width=1)
    draw.line((x + w - 16, y + 16, x + w, y + 16), fill=LINE, width=1)
    draw_center_text(draw, (x + 6, y + 6, x + w - 6, y + h - 6), text, ft(13))


def line(draw, pts, width=1):
    draw.line(pts, fill="black", width=width)


def crowfoot(draw, x, y, direction):
    if direction == "left":
        line(draw, [(x, y), (x + 12, y - 7)])
        line(draw, [(x, y), (x + 12, y)])
        line(draw, [(x, y), (x + 12, y + 7)])
    elif direction == "right":
        line(draw, [(x, y), (x - 12, y - 7)])
        line(draw, [(x, y), (x - 12, y)])
        line(draw, [(x, y), (x - 12, y + 7)])
    elif direction == "up":
        line(draw, [(x, y), (x - 7, y + 12)])
        line(draw, [(x, y), (x, y + 12)])
        line(draw, [(x, y), (x + 7, y + 12)])
    elif direction == "down":
        line(draw, [(x, y), (x - 7, y - 12)])
        line(draw, [(x, y), (x, y - 12)])
        line(draw, [(x, y), (x + 7, y - 12)])


def one_mark(draw, x, y, direction):
    if direction in ("left", "right"):
        draw.line((x, y - 8, x, y + 8), fill="black", width=1)
        offset = 5 if direction == "left" else -5
        draw.line((x + offset, y - 8, x + offset, y + 8), fill="black", width=1)
    else:
        draw.line((x - 8, y, x + 8, y), fill="black", width=1)
        offset = 5 if direction == "up" else -5
        draw.line((x - 8, y + offset, x + 8, y + offset), fill="black", width=1)


def zero_mark(draw, x, y):
    draw.ellipse((x - 5, y - 5, x + 5, y + 5), outline="black", width=1)


def label(draw, x, y, text):
    draw_center_text(draw, (x - 28, y - 14, x + 28, y + 14), text, ft(15))


def relation_horizontal(draw, x1, y1, x2, y2, *, left_card="one", right_card="many", text=None, text_xy=None):
    line(draw, [(x1, y1), (x2, y2)])
    if left_card == "one":
        one_mark(draw, x1, y1, "left")
    elif left_card == "zero-many":
        zero_mark(draw, x1 + 8, y1)
        crowfoot(draw, x1, y1, "left")
    elif left_card == "many":
        crowfoot(draw, x1, y1, "left")
    if right_card == "one":
        one_mark(draw, x2, y2, "right")
    elif right_card == "zero-many":
        zero_mark(draw, x2 - 8, y2)
        crowfoot(draw, x2, y2, "right")
    elif right_card == "many":
        crowfoot(draw, x2, y2, "right")
    if text and text_xy:
        label(draw, text_xy[0], text_xy[1], text)


def relation_poly(draw, points, text=None, text_xy=None):
    line(draw, points)
    if text and text_xy:
        label(draw, text_xy[0], text_xy[1], text)


def fig33():
    img = Image.new("RGB", (1180, 820), "white")
    draw = ImageDraw.Draw(img)
    draw_center_text(draw, (0, 10, 1180, 50), "核心交易模块数据库表关系图", ft(22, bold=True))

    category = draw_table(draw, 10, 620, 180, "商品分类表", "category", [
        "*分类编号：varchar(32)",
        "分类名称：varchar(100)",
        "分类图片：varchar(255)",
        "分类描述：varchar(255)",
        "创建时间：datetime",
    ])
    product = draw_table(draw, 300, 565, 260, "商品信息表", "product", [
        "*商品编号：varchar(32)",
        "所属分类编号：varchar(32)",
        "商品名称：varchar(255)",
        "商品描述：varchar(500)",
        "商品价格：decimal(10,2)",
        "商品图片：varchar(255)",
        "创建时间：datetime",
    ])
    user = draw_table(draw, 290, 320, 250, "用户表", "sys_user", [
        "*用户编号：bigint",
        "用户账号：varchar(30)",
        "用户昵称：varchar(30)",
        "手机号：varchar(11)",
        "头像地址：varchar(100)",
        "账号状态：char(1)",
        "创建时间：datetime",
    ])
    cart = draw_table(draw, 670, 600, 220, "购物车表", "cart", [
        "*购物车编号：varchar(32)",
        "商品编号：varchar(32)",
        "商品数量：int",
        "所属用户编号：bigint",
        "创建时间：datetime",
    ])
    address = draw_table(draw, 640, 360, 285, "收货地址表", "address", [
        "*地址编号：varchar(32)",
        "收货人：varchar(100)",
        "联系电话：varchar(20)",
        "详细地址：varchar(255)",
        "是否默认地址：tinyint(1)",
        "所属用户编号：bigint",
        "创建时间：datetime",
    ])
    order = draw_table(draw, 610, 45, 300, "订单购买关系", "order_purchase", [
        "*订单编号：varchar(32)",
        "下单用户编号：bigint",
        "商品编号：varchar(32)",
        "下单时间：datetime",
        "订单状态：varchar(200)",
        "订单总金额：decimal(10,2)",
        "购买数量：int",
        "成交价格：decimal(10,2)",
    ])

    relation_horizontal(draw, category["right"][0], category["right"][1], product["x"], category["right"][1], left_card="one", right_card="zero-many", text="包含", text_xy=(245, 685))
    relation_poly(draw, [(user["left"][0], user["left"][1] + 60), (205, user["left"][1] + 60), (205, address["left"][1]), (address["left"][0], address["left"][1])], text="拥有", text_xy=(215, 430))
    relation_poly(draw, [(user["left"][0], user["left"][1] + 120), (205, user["left"][1] + 120), (205, cart["left"][1]), (cart["left"][0], cart["left"][1])], text="拥有", text_xy=(215, 560))
    relation_horizontal(draw, product["right"][0], cart["x"], product["right"][1] + 30, product["right"][1] + 30, left_card="one", right_card="zero-many", text="加入", text_xy=(620, 690))

    relation_poly(draw, [(user["top"][0] + 10, user["y"]), (user["top"][0] + 10, 255), (order["x"], 255), (order["x"], order["y"] + 135)], text="下单", text_xy=(410, 250))
    one_mark(draw, user["top"][0] + 10, user["y"], "up")
    zero_mark(draw, order["x"] + 8, order["y"] + 135)
    crowfoot(draw, order["x"], order["y"] + 135, "left")

    relation_horizontal(draw, order["right"][0], order["right"][1] + 20, address["x"], order["right"][1] + 20, left_card="one", right_card="zero-many", text="关联地址", text_xy=(960, 245))
    relation_poly(draw, [(order["bottom"][0], order["bottom"][1]), (order["bottom"][0], 520), (product["top"][0], 520), (product["top"][0], product["y"])], text="购买商品", text_xy=(575, 470))
    one_mark(draw, order["bottom"][0], order["bottom"][1], "down")
    zero_mark(draw, product["top"][0], product["y"] - 8)
    crowfoot(draw, product["top"][0], product["y"], "up")

    draw_note(draw, 930, 135, 225, 72, "order_purchase 为兼容视图，\n由订单主表 order 与\n订单明细表 order_products\n组合形成。")

    img.save(FIG33)


def fig36():
    img = Image.new("RGB", (1100, 560), "white")
    draw = ImageDraw.Draw(img)
    draw_center_text(draw, (0, 10, 1100, 50), "算法推荐模块数据库表关系图", ft(22, bold=True))

    user = draw_table(draw, 30, 85, 240, "用户表", "sys_user", [
        "*用户编号：bigint",
        "用户账号：varchar(30)",
        "用户昵称：varchar(30)",
        "手机号：varchar(11)",
    ])
    category = draw_table(draw, 20, 330, 240, "商品分类表", "category", [
        "*分类编号：varchar(32)",
        "分类名称：varchar(100)",
        "分类图片：varchar(255)",
    ])
    product = draw_table(draw, 420, 350, 250, "商品表", "product", [
        "*商品编号：varchar(32)",
        "所属分类编号：varchar(32)",
        "商品名称：varchar(255)",
        "商品价格：decimal(10,2)",
        "商品图片：varchar(255)",
    ])
    order = draw_table(draw, 430, 75, 290, "订单购买关系", "order_purchase", [
        "*订单编号：varchar(32)",
        "下单用户编号：bigint",
        "商品编号：varchar(32)",
        "下单时间：datetime",
        "订单状态：varchar(200)",
        "订单总金额：decimal(10,2)",
        "购买数量：int",
        "成交价格：decimal(10,2)",
    ])

    relation_horizontal(draw, user["right"][0], order["x"], user["right"][1] + 20, user["right"][1] + 60, left_card="one", right_card="zero-many", text="下单", text_xy=(350, 205))
    relation_horizontal(draw, category["right"][0], product["x"], category["right"][1] + 40, category["right"][1] + 40, left_card="one", right_card="zero-many", text="包含", text_xy=(340, 415))
    relation_poly(draw, [(order["bottom"][0], order["bottom"][1]), (order["bottom"][0], 315), (product["top"][0], 315), (product["top"][0], product["y"])], text="购买", text_xy=(580, 315))
    one_mark(draw, order["bottom"][0], order["bottom"][1], "down")
    zero_mark(draw, product["top"][0], product["y"] - 8)
    crowfoot(draw, product["top"][0], product["y"], "up")

    draw_note(draw, 780, 130, 290, 78, "推荐模块通过读取 order_purchase\n中的历史购买记录构造用户—商品\n交互关系，并结合商品信息完成\n候选结果计算。")
    img.save(FIG36)


def set_doc_font(run, size=12, east_asia="宋体", western="Times New Roman", bold=False):
    run.bold = bold
    run.font.name = western
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_doc_paragraph(doc, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, east_asia="宋体"):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = Pt(20)
    r = p.add_run(text)
    set_doc_font(r, size=size, east_asia=east_asia, bold=bold)
    return p


def build_docx():
    doc = Document()
    add_doc_paragraph(doc, "图3-3与图3-6-更新版", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, east_asia="黑体")
    add_doc_paragraph(doc, "图3-3 核心交易模块数据库表关系图（更新版）", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, east_asia="黑体")
    doc.add_picture(str(FIG33), width=Inches(6.8))
    add_doc_paragraph(doc, "图3-3 核心交易模块数据库表关系图（更新版）", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, east_asia="黑体")
    add_doc_paragraph(doc, "图3-6 算法推荐模块数据库表关系图（更新版）", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, east_asia="黑体")
    doc.add_picture(str(FIG36), width=Inches(6.8))
    add_doc_paragraph(doc, "图3-6 算法推荐模块数据库表关系图（更新版）", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, east_asia="黑体")
    doc.save(DOCX_OUT)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    fig33()
    fig36()
    build_docx()
    print(FIG33)
    print(FIG36)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()
