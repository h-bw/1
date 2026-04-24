from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


OUTPUT = r"D:\DevelopSoft\lingshi\3.2.4数据库表设计-一致性修改稿.docx"


def set_run_font(run, size=12, east_asia="宋体", western="Times New Roman", bold=False):
    run.bold = bold
    run.font.name = western
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=None):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if line is not None:
        fmt.line_spacing = Pt(line)


def set_three_line_table(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = borders.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            borders.append(element)
        if name in {"top", "bottom", "insideH"}:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "8")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")
        else:
            element.set(qn("w:val"), "nil")


def add_text(doc, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, east_asia="宋体", before=0, after=0, line=20):
    p = doc.add_paragraph()
    set_paragraph(p, align=align, before=before, after=after, line=line)
    run = p.add_run(text)
    set_run_font(run, size=size, east_asia=east_asia, bold=bold)
    return p


def add_table_title(doc, text):
    add_text(
        doc,
        text,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=10,
        east_asia="黑体",
        before=12,
        after=0,
        line=20,
    )


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_three_line_table(table)

    for i, value in enumerate(headers):
        p = table.rows[0].cells[i].paragraphs[0]
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(value)
        set_run_font(run, size=10, east_asia="宋体", bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
            run = p.add_run(value)
            set_run_font(run, size=10, east_asia="宋体")


def main():
    doc = Document()

    add_text(
        doc,
        "3.2.4 数据库表设计",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        size=16,
        east_asia="黑体",
        before=0,
        after=6,
        line=20,
    )

    add_text(
        doc,
        "为与前文概念结构设计和逻辑结构设计保持一致，数据库表设计围绕用户、收货地址、商品分类、商品信息、购物车和订单购买六类核心对象展开。结合系统当前实现，用户信息由 sys_user 表承担，订单购买关系则由订单主表、订单明细表组合形成，并在数据库中通过 order_purchase 兼容视图进行表示。这样既能保持论文中 E-R 图、关系模式与表设计的统一，也不会破坏现有系统的订单处理流程。",
        size=12,
    )

    add_text(
        doc,
        "用户表用于存储系统用户的基本信息，为登录认证、订单归属和购物车管理提供用户数据支撑，其结构如表3-2所示。",
        size=12,
    )
    add_table_title(doc, "表3-2 用户表")
    add_table(
        doc,
        ["字段名", "类型", "主键", "非空", "说明"],
        [
            ["user_id", "bigint", "是", "是", "用户编号"],
            ["user_name", "varchar(30)", "否", "是", "用户名"],
            ["phonenumber", "varchar(11)", "否", "否", "手机号"],
            ["create_time", "datetime", "否", "否", "创建时间"],
        ],
    )
    add_text(
        doc,
        "其中，user_id 是主键，用于唯一标识系统中的每一位用户；user_name 用于存储登录账号；phonenumber 用于存储联系方式；create_time 用于记录用户创建时间。该表在系统中主要用于标识订单归属用户及购物车所属用户。",
        size=12,
    )

    add_text(
        doc,
        "收货地址表用于存储用户的收货地址信息，便于前端下单时选择地址以及后端进行订单联系信息核验，其结构如表3-3所示。",
        size=12,
    )
    add_table_title(doc, "表3-3 收货地址表")
    add_table(
        doc,
        ["字段名", "类型", "主键", "非空", "说明"],
        [
            ["address_id", "varchar(32)", "是", "是", "地址编号"],
            ["user_id", "bigint", "否", "是", "所属用户编号"],
            ["name", "varchar(100)", "否", "是", "收货人"],
            ["phone", "varchar(20)", "否", "是", "联系电话"],
            ["detail", "varchar(255)", "否", "是", "详细地址"],
            ["is_default", "tinyint(1)", "否", "是", "是否默认地址"],
            ["create_time", "datetime", "否", "是", "创建时间"],
        ],
    )
    add_text(
        doc,
        "其中，address_id 是主键，用于唯一标识一条收货地址记录；user_id 为外键，用于关联用户表；detail 用于存储详细地址信息；is_default 用于标识该地址是否为默认地址。该表与用户表之间构成一对多关系。",
        size=12,
    )

    add_text(
        doc,
        "商品分类表用于存储商城中的零食商品分类信息，供前台分类浏览以及后台分类管理使用，其结构如表3-4所示。",
        size=12,
    )
    add_table_title(doc, "表3-4 商品分类表")
    add_table(
        doc,
        ["字段名", "类型", "主键", "非空", "说明"],
        [
            ["category_id", "varchar(32)", "是", "是", "分类编号"],
            ["name", "varchar(100)", "否", "是", "分类名称"],
            ["image", "varchar(255)", "否", "是", "分类图片"],
            ["description", "varchar(255)", "否", "是", "分类描述"],
            ["create_time", "datetime", "否", "否", "创建时间"],
        ],
    )
    add_text(
        doc,
        "其中，category_id 是主键，用于标识每个商品分类；name 用于存储分类名称；image 用于存储分类展示图片；description 用于补充分类描述。该表与商品信息表之间通过分类编号建立联系。",
        size=12,
    )

    add_text(
        doc,
        "商品信息表用于存储平台上的零食商品基本信息，是前端商品展示、分类检索和后台商品管理的核心数据表，其结构如表3-5所示。",
        size=12,
    )
    add_table_title(doc, "表3-5 商品信息表")
    add_table(
        doc,
        ["字段名", "类型", "主键", "非空", "说明"],
        [
            ["product_id", "varchar(32)", "是", "是", "商品编号"],
            ["category_id", "varchar(32)", "否", "是", "所属分类编号"],
            ["name", "varchar(255)", "否", "是", "商品名称"],
            ["description", "varchar(500)", "否", "是", "商品描述"],
            ["price", "decimal(10,2)", "否", "是", "商品价格"],
            ["image", "varchar(255)", "否", "是", "商品图片"],
            ["create_time", "datetime", "否", "否", "创建时间"],
        ],
    )
    add_text(
        doc,
        "其中，product_id 是主键，用于唯一标识一个商品；category_id 为外键，用于连接商品分类；name、description、price 和 image 分别用于存储商品名称、描述、价格和图片信息。该表支撑商品浏览、商品详情查询和后台商品维护等功能。",
        size=12,
    )

    add_text(
        doc,
        "购物车表用于存储用户在生成订单之前所选商品的信息，是商品浏览到订单提交之间的桥梁数据表，其结构如表3-6所示。",
        size=12,
    )
    add_table_title(doc, "表3-6 购物车表")
    add_table(
        doc,
        ["字段名", "类型", "主键", "非空", "说明"],
        [
            ["cart_id", "varchar(32)", "是", "是", "购物车编号"],
            ["user_id", "bigint", "否", "是", "所属用户编号"],
            ["product_id", "varchar(32)", "否", "是", "商品编号"],
            ["quantity", "int", "否", "是", "商品数量"],
            ["create_time", "datetime", "否", "是", "创建时间"],
        ],
    )
    add_text(
        doc,
        "其中，cart_id 是主键，用于唯一标识一条购物车记录；user_id 和 product_id 分别关联用户与商品；quantity 用于记录当前加购数量。该表反映了用户与商品之间的暂存购买关系，可用于查询购物车商品、修改数量和结算生成订单。",
        size=12,
    )

    add_text(
        doc,
        "订单购买表用于表示用户与商品之间在下单场景下形成的购买关系，其结构如表3-7所示。结合系统当前实现，该表由订单主信息与订单商品信息组合形成，并通过 order_purchase 兼容视图进行表示，以便与前文逻辑结构设计中的“订单购买”关系保持一致。",
        size=12,
    )
    add_table_title(doc, "表3-7 订单购买表")
    add_table(
        doc,
        ["字段名", "类型", "主键", "非空", "说明"],
        [
            ["order_id", "varchar(32)", "是", "是", "订单编号"],
            ["user_id", "bigint", "否", "是", "下单用户编号"],
            ["product_id", "varchar(32)", "否", "否", "商品编号"],
            ["create_time", "datetime", "否", "否", "下单时间"],
            ["status", "varchar(200)", "否", "是", "订单状态"],
            ["total_amount", "decimal(10,2)", "否", "是", "订单总金额"],
            ["quantity", "int", "否", "是", "购买数量"],
            ["price", "decimal(10,2)", "否", "是", "成交价格"],
        ],
    )
    add_text(
        doc,
        "其中，order_id 用于标识订单记录，user_id 用于关联下单用户，product_id 用于标识订单中的商品，quantity 和 price 分别表示购买数量与成交价格，status 和 total_amount 分别表示订单状态和订单总金额。该表体现了用户与商品之间的下单购买联系，也为订单查询、推荐分析和后续数据统计提供了基础。",
        size=12,
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
