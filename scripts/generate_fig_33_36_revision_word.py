from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


OUTPUT = r"D:\DevelopSoft\lingshi\图3-3与图3-6修改内容.docx"


def set_run_font(run, size=12, east_asia="宋体", western="Times New Roman", bold=False):
    run.bold = bold
    run.font.name = western
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=20):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = Pt(line)


def add_para(doc, text, *, bold=False, size=12, east_asia="宋体", align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0):
    p = doc.add_paragraph()
    set_paragraph(p, align=align, before=before, after=after)
    r = p.add_run(text)
    set_run_font(r, size=size, east_asia=east_asia, bold=bold)
    return p


def main():
    doc = Document()

    add_para(
        doc,
        "图3-3与图3-6修改内容",
        bold=True,
        size=16,
        east_asia="黑体",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=6,
    )
    add_para(
        doc,
        "以下内容仅包含图3-3和图3-6需要替换的图名、图中节点命名建议及正文说明，可直接用于统一数据库章节中“订单购买关系”的表达口径。",
        size=12,
    )

    add_para(doc, "一、图3-3 修改建议", bold=True, size=14, east_asia="黑体", align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=3)
    add_para(
        doc,
        "图中节点命名建议：将原图中的“订单主表”“订单明细表”统一替换为“订单购买（订单主表+订单明细表）”。如果图中空间有限，也可以直接写成“订单购买”。",
        size=12,
    )
    add_para(
        doc,
        "图名建议：图3-3 核心交易模块数据库表关系图",
        size=12,
    )
    add_para(
        doc,
        "正文说明替换稿：图3-3 主要描述系统核心交易流程中各数据对象之间的关系，包括用户表、收货地址表、商品信息表、购物车表以及订单购买关系等。其中，用户与商品之间通过购物车关系和订单购买关系建立联系，商品信息则与分类信息共同支撑商品展示和交易处理。在系统实际实现中，订单购买关系由订单主表与订单明细表共同支撑，因此该图既反映了交易流程中的主要数据联系，也体现了系统在数据库层面的核心结构。",
        size=12,
    )

    add_para(doc, "二、图3-6 修改建议", bold=True, size=14, east_asia="黑体", align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=3)
    add_para(
        doc,
        "图中节点命名建议：将原图中的“订单主表”“订单明细表”统一替换为“订单购买（订单主表+订单明细表）”。如果图中更强调逻辑抽象层，也可以直接写成“订单购买”。",
        size=12,
    )
    add_para(
        doc,
        "图名建议：图3-6 算法推荐模块数据库表关系图",
        size=12,
    )
    add_para(
        doc,
        "正文说明替换稿：图3-6 展示了推荐模块所依赖的核心数据对象及其联系，包括用户表、订单购买关系、商品信息表和商品分类表等。推荐模块通过读取用户历史购买记录，构造用户—商品交互关系，并在此基础上计算候选推荐结果。在系统实际实现中，订单购买关系的数据来源由订单主表与订单明细表共同提供，因此该图在逻辑层面上统一表达为“订单购买关系参与推荐计算”，在实现层面上则对应订单主从表结构提供的数据支撑。",
        size=12,
    )

    add_para(doc, "三、统一口径建议", bold=True, size=14, east_asia="黑体", align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=3)
    add_para(
        doc,
        "为保证第三章数据库相关内容前后一致，建议全文统一采用以下表达方式：在概念层和逻辑层中使用“订单购买关系”这一表述；在系统实现层中说明该关系由订单主表与订单明细表共同实现；在数据库兼容表示层中，可补充说明该关系通过 order_purchase 视图进行统一组织。这样既能与 E-R 图、关系模式和数据库表设计保持一致，也能与系统当前真实实现相吻合。",
        size=12,
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
