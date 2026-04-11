from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SRC = Path(r"D:\DevelopSoft\lingshi\论文Word排版版-可直接贴入.md")
DST = Path(r"D:\DevelopSoft\lingshi\基于SpringBoot的助农扶贫商城设计与实现-精修版.docx")


def set_doc_defaults(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)


def set_run_style(run, east_asia="宋体", latin="Times New Roman", size=12, bold=False):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_format(p, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.5, before=0, after=0):
    p.alignment = align
    fmt = p.paragraph_format
    fmt.line_spacing = line
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if first_line:
        fmt.first_line_indent = Cm(0.74)
    else:
        fmt.first_line_indent = Cm(0)


def add_paragraph(doc, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, east_asia="宋体"):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line, align=align)
    r = p.add_run(text)
    set_run_style(r, east_asia=east_asia, size=size, bold=bold)
    return p


def add_heading(doc, text, level):
    if level == 1:
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6)
        r = p.add_run(text)
        set_run_style(r, east_asia="黑体", size=16, bold=True)
        return p
    if level == 2:
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=3)
        r = p.add_run(text)
        set_run_style(r, east_asia="黑体", size=14, bold=True)
        return p
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=3, after=3)
    r = p.add_run(text)
    set_run_style(r, east_asia="黑体", size=12, bold=True)
    return p


def add_cover_page(doc):
    for _ in range(4):
        doc.add_paragraph()
    add_paragraph(doc, "毕业设计（论文）", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, east_asia="黑体")
    for _ in range(4):
        doc.add_paragraph()
    add_paragraph(doc, "基于SpringBoot的助农扶贫商城设计与实现", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, east_asia="黑体")
    for _ in range(6):
        doc.add_paragraph()

    items = [
        "学生姓名：黄博文",
        "学号：220501050035",
        "学部（系）：信息科学与技术学部",
        "专业名称：软件工程",
        "指导教师：章勤",
        "职称或学位：教授",
        "日期：2026年5月31日",
    ]
    for item in items:
        add_paragraph(doc, item, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    doc.add_page_break()


def add_statement_page(doc, title, body):
    add_heading(doc, title, 1)
    add_paragraph(doc, body)
    doc.add_paragraph()
    add_paragraph(doc, "论文作者签名：_______________", first_line=False)
    add_paragraph(doc, "日期：_______________", first_line=False)
    doc.add_page_break()


def add_toc_placeholder(doc):
    add_heading(doc, "目 录", 1)
    p = add_paragraph(doc, "请在 Word 中使用“引用-目录-自动目录”生成目录。", first_line=False)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "摘要",
        "Abstract",
        "第一章 绪论",
        "第二章 系统设计结构及相关技术介绍",
        "第三章 系统设计与关键模块实现",
        "第四章 系统实现与测试",
        "第五章 总结与展望",
        "结束语",
        "致谢",
        "参考文献",
    ]:
        add_paragraph(doc, line, first_line=False)
    doc.add_page_break()


def is_table_line(line: str):
    return line.strip().startswith("|") and line.strip().endswith("|")


def parse_table(lines, start_idx):
    table_lines = []
    idx = start_idx
    while idx < len(lines) and is_table_line(lines[idx]):
        table_lines.append(lines[idx].strip())
        idx += 1
    return table_lines, idx


def split_md_row(line):
    parts = [p.strip() for p in line.strip().strip("|").split("|")]
    return parts


def add_markdown_table(doc, table_lines):
    rows = [split_md_row(line) for line in table_lines if not re.fullmatch(r"\|\s*[-: ]+\s*(\|\s*[-: ]+\s*)+\|", line)]
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run_style(r, size=10.5, bold=(i == 0))
            p.paragraph_format.line_spacing = 1.2


def add_numbered_paragraph(doc, line):
    p = doc.add_paragraph(style=None)
    set_paragraph_format(p, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    r = p.add_run(line)
    set_run_style(r, size=12)


def add_reference_paragraph(doc, line):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    fmt = p.paragraph_format
    fmt.left_indent = Cm(0.74)
    fmt.hanging_indent = Cm(0.74)
    r = p.add_run(line)
    set_run_style(r, size=12)


def build_body_from_markdown(doc):
    content = SRC.read_text(encoding="utf-8").splitlines()

    # Skip handmade front matter in markdown because we generate refined pages ourselves.
    start_idx = 0
    for idx, line in enumerate(content):
        if line.strip() == "# 基于SpringBoot的助农扶贫商城设计与实现":
            start_idx = idx
            break
    lines = content[start_idx:]

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        if line.strip() == "---":
            doc.add_page_break()
            i += 1
            continue

        if is_table_line(line):
            table_lines, i = parse_table(lines, i)
            add_markdown_table(doc, table_lines)
            continue

        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
            i += 1
            continue

        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
            i += 1
            continue

        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
            i += 1
            continue

        if re.match(r"^\d+\.\s", line.strip()):
            add_numbered_paragraph(doc, line.strip())
            i += 1
            continue

        if re.match(r"^\[\d+\]", line.strip()):
            add_reference_paragraph(doc, line.strip())
            i += 1
            continue

        if line.startswith("图") or line.startswith("表"):
            p = add_paragraph(doc, line, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        if line.startswith("此处插入图") or line.startswith("此处插入表") or line.startswith("说明："):
            p = add_paragraph(doc, line, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
            i += 1
            continue

        if line.startswith("关键词：") or line.startswith("Keywords:"):
            add_paragraph(doc, line, size=12, first_line=False)
            i += 1
            continue

        add_paragraph(doc, line)
        i += 1


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def build_doc():
    doc = Document()
    set_doc_defaults(doc)

    add_cover_page(doc)
    add_statement_page(
        doc,
        "毕业设计（论文）独创性声明",
        "本人声明所呈交的毕业设计（论文）是本人在导师指导下进行的研究工作及取得的研究成果。除了文中特别加以标注和致谢的地方外，本论文中不包含他人已经发表或撰写过的研究成果，也不包含为获得其他教育机构的学位或证书而使用过的材料。对本研究做出重要贡献的个人和集体，均已在文中作了明确说明并表示谢意。",
    )
    add_statement_page(
        doc,
        "毕业设计（论文）使用授权声明",
        "本人完全了解学校关于毕业设计（论文）保存、使用和管理的相关规定，同意学校保留并向有关部门送交论文的复印件和电子文档，允许论文被查阅和借阅；本人授权学校可以将论文的全部或部分内容编入有关数据库进行检索，并可采用影印、缩印或其他复制手段保存、汇编本论文。",
    )
    add_toc_placeholder(doc)
    build_body_from_markdown(doc)

    for section in doc.sections:
        add_page_number(section)

    doc.save(DST)


if __name__ == "__main__":
    build_doc()
    print(str(DST))
