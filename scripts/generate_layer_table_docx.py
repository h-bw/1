from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


OUT = Path(r"D:\DevelopSoft\lingshi\系统分层结构说明表-论文版.docx")


def set_font(run, size=12, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold


def set_paragraph(paragraph, text, size=12, bold=False, center=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)
    paragraph.paragraph_format.first_line_indent = Pt(24) if not center else Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5


def style_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=11, bold=bold)


def main():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    set_paragraph(
        doc.add_paragraph(),
        "在系统总体架构设计中，各层之间分工明确、职责清晰。为了更直观地说明本系统各层的组成及功能，现将系统分层结构说明如表2-1所示。",
        size=12,
    )

    set_paragraph(doc.add_paragraph(), "表2-1 系统分层结构说明表", size=12, bold=True, center=True)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["层次名称", "主要组成", "主要职责"]
    for i, text in enumerate(headers):
        style_cell(table.rows[0].cells[i], text, bold=True)

    rows = [
        ("前端展示层", "UniApp 用户端、Vue 后台管理端", "负责页面展示、用户交互、数据请求发起以及结果展示"),
        ("控制层", "Controller", "负责接收前端请求、参数封装、请求分发以及统一返回处理结果"),
        ("业务层", "Service", "负责实现商品查询、购物车操作、订单生成、地址管理、轮播图管理和推荐逻辑等具体业务处理"),
        ("持久层", "Mapper、MyBatis", "负责与数据库进行交互，完成各类业务数据的增删改查操作"),
        ("数据层", "MySQL 数据库", "负责存储系统中的用户、商品、分类、购物车、订单、地址和轮播图等核心数据"),
    ]

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            style_cell(cells[i], text, bold=False)

    set_paragraph(
        doc.add_paragraph(),
        "由表2-1可以看出，本系统采用了较为典型的分层架构设计。前端展示层主要负责页面呈现和用户操作，控制层负责请求接收与分发，业务层负责核心逻辑处理，持久层负责数据访问，数据层负责数据存储。通过这种分层方式，系统能够较好地降低模块之间的耦合度，使各部分职责更加明确，也为后续功能扩展与维护提供了便利。",
        size=12,
    )

    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    main()
