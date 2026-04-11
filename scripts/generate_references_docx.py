from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


OUT = Path(r"D:\DevelopSoft\lingshi\参考文献-论文版.docx")

REFERENCES = [
    "[1] 吴迁. 基于uni-app与Spring Boot框架的Web应用平台的设计与实现[D]. 西安: 西安石油大学, 2025.",
    "[2] Vue.js前端开发框架应用[M]. 北京: 人民邮电出版社, 2024.",
    "[3] 十三. Spring Cloud Alibaba大型微服务架构项目实战[M]. 北京: 电子工业出版社, 2024.",
    "[4] Spring Cloud Alibaba微服务框架电商平台搭建与编程解析[M]. 北京: 人民邮电出版社, 2023.",
    "[5] Spring+Spring MVC+MyBatis整合开发实战[M]. 北京: 机械工业出版社, 2022.",
    "[6] 牛子逸. 基于Vue+SpringBoot的音乐评阅系统设计与实现[D]. 成都: 电子科技大学, 2025. DOI:10.27005/d.cnki.gdzku.2025.003877.",
    "[7] 马艳夕. 基于SpringBoot与Vue技术的企业电商平台的设计与实现[J]. 信息与电脑(理论版), 2021, 33(3): 99-100.",
    "[8] 陈小燕, 朱映辉, 余晓春. 基于SpringBoot+Vue的好农物商城的设计与实现[J]. 电脑知识与技术, 2022, 18(22): 37-39.",
    "[9] 梅忠. 基于UniApp与Spring Boot的校友平台设计与实现[J]. 信息记录材料, 2025, 26(11): 113-116.",
    "[10] 李宜镓. 基于SpringBoot的电商秒杀系统的设计与实现[D]. 西安: 西安电子科技大学, 2022.",
    "[11] Yang N. Personal health information service platform based on Vue.js+SpringBoot[J]. The Frontiers of Society, Science and Technology, 2025(5): 510-517.",
    "[12] Cuomo S. Vue.js 3 for beginners: learn the essentials of Vue.js 3 and its ecosystem to build modern web applications[M]. 2024.",
    "[13] Obaid A J, Schiopoiu A B, Bhushan B, et al. E-commerce, marketing, and consumer behavior in the AI era[M]. 2024.",
    "[14] Al-Ma’aitah A M. Critical success factors for e-commerce: a literature review[J]. International Journal of Advanced Computer Science and Applications, 2020, 11(9): 375-381.",
    "[15] Wu H J. Commerce middle office management system based on SpringBoot[J]. International Journal of Advanced Network, Monitoring and Controls, 2022(2): 32-45.",
]


def set_font(run, size=12, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold


def main():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.5
    run = title.add_run("参考文献")
    set_font(run, size=14, bold=True)

    for ref in REFERENCES:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(ref)
        set_font(run, size=12)

    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    main()
