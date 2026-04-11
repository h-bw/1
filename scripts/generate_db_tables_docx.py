from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from pathlib import Path


OUT = Path(r"D:\DevelopSoft\lingshi\数据库表设计-论文版.docx")


TABLES = [
    {
        "title": "表3-2 用户信息表 sys_user",
        "rows": [
            ["字段名", "类型", "主键", "非空", "说明"],
            ["user_id", "bigint", "是", "是", "用户编号"],
            ["user_name", "varchar(30)", "否", "是", "用户账号"],
            ["nick_name", "varchar(30)", "否", "是", "用户昵称"],
            ["phonenumber", "varchar(11)", "否", "否", "手机号码"],
            ["avatar", "varchar(100)", "否", "否", "头像地址"],
            ["password", "varchar(100)", "否", "否", "登录密码"],
            ["status", "char(1)", "否", "否", "账号状态"],
            ["create_time", "datetime", "否", "否", "创建时间"],
        ],
    },
    {
        "title": "表3-3 商品分类表 category",
        "rows": [
            ["字段名", "类型", "主键", "非空", "说明"],
            ["category_id", "varchar(255)", "是", "是", "分类编号"],
            ["name", "varchar(100)", "否", "是", "分类名称"],
            ["image", "varchar(255)", "否", "是", "分类图片"],
            ["description", "varchar(255)", "否", "否", "分类描述"],
            ["create_time", "datetime", "否", "否", "创建时间"],
        ],
    },
    {
        "title": "表3-4 商品信息表 product",
        "rows": [
            ["字段名", "类型", "主键", "非空", "说明"],
            ["product_id", "varchar(255)", "是", "是", "商品编号"],
            ["category_id", "varchar(255)", "否", "是", "所属分类编号"],
            ["name", "varchar(255)", "否", "是", "商品名称"],
            ["description", "text", "否", "否", "商品描述"],
            ["price", "decimal(10,2)", "否", "是", "商品价格"],
            ["image", "varchar(255)", "否", "否", "商品图片"],
            ["create_time", "datetime", "否", "否", "创建时间"],
        ],
    },
    {
        "title": "表3-5 购物车表 cart",
        "rows": [
            ["字段名", "类型", "主键", "非空", "说明"],
            ["cart_id", "varchar(255)", "是", "是", "购物车编号"],
            ["product_id", "varchar(255)", "否", "是", "商品编号"],
            ["quantity", "int", "否", "是", "商品数量"],
            ["user_id", "bigint", "否", "是", "所属用户编号"],
            ["create_time", "datetime", "否", "否", "创建时间"],
        ],
    },
    {
        "title": "表3-6 收货地址表 address",
        "rows": [
            ["字段名", "类型", "主键", "非空", "说明"],
            ["address_id", "varchar(255)", "是", "是", "地址编号"],
            ["name", "varchar(100)", "否", "是", "收货人姓名"],
            ["phone", "varchar(20)", "否", "是", "联系电话"],
            ["detail", "text", "否", "是", "详细地址"],
            ["is_default", "tinyint(1)", "否", "否", "是否默认地址"],
            ["user_id", "bigint", "否", "是", "所属用户编号"],
            ["create_time", "datetime", "否", "否", "创建时间"],
        ],
    },
    {
        "title": "表3-7 订单主表 order",
        "rows": [
            ["字段名", "类型", "主键", "非空", "说明"],
            ["order_id", "varchar(255)", "是", "是", "订单编号"],
            ["name", "varchar(255)", "否", "是", "收货人姓名"],
            ["phone", "varchar(255)", "否", "是", "联系电话"],
            ["address", "varchar(255)", "否", "是", "收货地址"],
            ["total_amount", "decimal(10,2)", "否", "是", "订单总金额"],
            ["product_count", "int", "否", "是", "商品总件数"],
            ["status", "varchar(200)", "否", "是", "订单状态"],
            ["remark", "text", "否", "否", "订单备注"],
            ["user_id", "bigint", "否", "是", "下单用户编号"],
            ["create_time", "datetime", "否", "否", "下单时间"],
        ],
    },
    {
        "title": "表3-8 订单明细表 order_products",
        "rows": [
            ["字段名", "类型", "主键", "非空", "说明"],
            ["op_id", "bigint", "是", "是", "订单明细编号"],
            ["order_id", "varchar(255)", "否", "是", "所属订单编号"],
            ["image", "varchar(255)", "否", "否", "商品图片"],
            ["name", "varchar(255)", "否", "是", "商品名称"],
            ["price", "decimal(10,2)", "否", "是", "商品价格"],
            ["quantity", "int", "否", "是", "商品数量"],
        ],
    },
    {
        "title": "表3-9 轮播图表 banner",
        "rows": [
            ["字段名", "类型", "主键", "非空", "说明"],
            ["banner_id", "varchar(255)", "是", "是", "轮播图编号"],
            ["image", "varchar(255)", "否", "是", "图片地址"],
            ["sort", "int", "否", "是", "排序值"],
        ],
    },
]


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size=12, bold=False):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def style_normal(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, east_asia="黑体", size=12, bold=True)


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            set_run_font(r, size=10.5, bold=(i == 0))


def main():
    doc = Document()
    style_normal(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("数据库表设计")
    set_run_font(r, east_asia="黑体", size=16, bold=True)

    intro = doc.add_paragraph()
    intro.paragraph_format.first_line_indent = Cm(0.74)
    r = intro.add_run("以下内容根据当前零食商城小程序项目整理形成论文可用的数据库表设计格式，可直接插入第三章数据库设计部分。")
    set_run_font(r, size=12)

    for item in TABLES:
        doc.add_paragraph()
        add_title(doc, item["title"])
        add_table(doc, item["rows"])

    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    main()
