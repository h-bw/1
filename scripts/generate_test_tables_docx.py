from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


OUT = Path(r"D:\DevelopSoft\lingshi\系统测试表格-论文版.docx")


def set_font(run, size=12, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold


def add_paragraph(doc, text, size=12, bold=False, center=False, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Pt(24) if (indent and not center) else Pt(0)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p


def style_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=10.5, bold=bold)


def add_test_table(doc, title, rows):
    add_paragraph(doc, title, size=12, bold=True, center=True, indent=False)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["测试编号", "测试内容", "测试步骤", "预期结果", "测试结果"]
    for i, h in enumerate(headers):
        style_cell(table.rows[0].cells[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            style_cell(cells[i], text)
    doc.add_paragraph()


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    add_paragraph(doc, "4.5 系统测试", size=14, bold=True, center=False, indent=False)

    add_paragraph(
        doc,
        "为检验零食商城小程序系统的功能是否完善、运行是否稳定，本文对系统进行了功能测试。测试内容主要围绕用户端和后台管理端的关键业务流程展开，重点包括首页展示、商品浏览、购物车操作、订单提交、订单查询，以及后台商品管理、后台订单管理和推荐功能展示等内容。",
    )

    add_paragraph(doc, "4.5.1 测试方式", size=12, bold=True, center=False, indent=False)
    add_paragraph(
        doc,
        "本系统主要采用功能测试的方法，对各模块在实际使用过程中的输入、处理过程和输出结果进行检查。测试过程中，重点关注前后端接口能否正常返回数据，页面能否准确显示相关内容，数据库记录是否与页面操作保持一致，以及系统在多次连续操作情况下是否仍能稳定运行。",
    )

    add_paragraph(doc, "根据系统的主要业务流程，本文选取了具有代表性的测试场景进行验证，测试结果分别如表4-1至表4-4所示。")

    add_test_table(doc, "表4-1 用户端主要功能测试表", [
        ("T01", "首页展示", "进入系统首页，查看轮播图和推荐商品信息", "首页内容正常加载，轮播图和商品列表显示正确", "通过"),
        ("T02", "商品分类浏览", "进入分类页面，点击不同分类查看商品", "能够正确显示对应分类下的商品信息", "通过"),
        ("T03", "商品详情查看", "点击商品进入详情页面", "商品名称、价格、图片和描述等信息显示正确", "通过"),
        ("T04", "加入购物车", "在商品详情页点击加入购物车", "购物车中新增对应商品，数量更新正确", "通过"),
        ("T05", "购物车数量修改", "在购物车页面增加或减少商品数量", "购物车页面和数据库中的商品数量保持一致", "通过"),
        ("T06", "收货地址管理", "新增、修改并选择收货地址", "地址信息保存成功，可在下单时正常选择", "通过"),
    ])

    add_test_table(doc, "表4-2 订单模块功能测试表", [
        ("T07", "订单确认", "在购物车中选择商品并进入订单确认页面", "订单页面能够正确显示商品、地址和金额信息", "通过"),
        ("T08", "提交订单", "确认订单信息后提交订单", "订单主表和订单明细表写入成功，返回下单成功结果", "通过"),
        ("T09", "购物车数据清理", "订单提交成功后返回购物车页面", "已结算商品从购物车中删除", "通过"),
        ("T10", "我的订单查询", "进入我的订单页面查看历史订单", "能够正常显示订单编号、金额、状态和明细信息", "通过"),
    ])

    add_test_table(doc, "表4-3 后台管理功能测试表", [
        ("T11", "商品分类管理", "在后台新增、修改和删除商品分类", "分类信息可正常维护并实时生效", "通过"),
        ("T12", "商品管理", "在后台新增、修改和删除商品信息", "商品数据保存正确，前台可同步显示", "通过"),
        ("T13", "轮播图管理", "在后台维护轮播图信息", "轮播图新增、修改和删除操作正常", "通过"),
        ("T14", "订单管理", "在后台查看订单列表和订单详情", "订单编号、金额、状态等信息显示正确", "通过"),
    ])

    add_test_table(doc, "表4-4 推荐功能测试表", [
        ("T15", "热门商品推荐", "使用无历史订单的用户访问推荐内容", "系统能够返回热门商品列表", "通过"),
        ("T16", "历史订单推荐", "使用存在历史订单的用户访问推荐内容", "系统能够根据历史订单返回推荐商品", "通过"),
        ("T17", "推荐结果展示", "在前端页面查看推荐商品信息", "推荐商品列表能够正常显示", "通过"),
    ])

    add_paragraph(doc, "4.5.2 测试结果分析", size=12, bold=True, center=False, indent=False)
    add_paragraph(
        doc,
        "测试结果表明，系统首页展示、商品浏览、购物车管理、订单提交以及后台管理等主要功能均可正常使用。用户在商品详情页面将商品加入购物车后，购物车列表能够及时更新相应内容；订单提交后，订单主表和订单明细表中的数据能够正常保存，已结算的购物车商品也会同步删除。后台管理员能够对商品分类、商品信息、轮播图以及订单信息进行维护。推荐功能也可以根据用户历史订单或热门商品统计结果返回相应的商品列表。",
    )
    add_paragraph(
        doc,
        "从测试情况来看，系统在当前设计范围内已经能够完成零食商城小程序的基本业务流程，各模块之间衔接较为顺畅，运行结果与预期基本一致。",
    )

    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    main()
