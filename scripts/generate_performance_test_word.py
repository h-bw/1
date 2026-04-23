import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"D:\DevelopSoft\lingshi")
JSON_PATH = ROOT / "logs" / "interface_benchmark_results.json"
OUTPUT_PATH = ROOT / "4.3性能测试结果-Word版.docx"


def set_font(run, east_asia: str = "宋体", western: str = "Times New Roman", size: int = 10):
    run.font.name = western
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if line is not None:
        fmt.line_spacing = Pt(line)


def add_text_paragraph(doc: Document, text: str, *, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       east_asia="宋体", western="Times New Roman", size=10, before=0, after=0):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    set_font(run, east_asia=east_asia, western=western, size=size)
    set_paragraph_spacing(p, before=before, after=after)
    return p


def set_cell_text(cell, text: str, *, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    set_font(run, size=10)


def set_three_line_table(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = tbl_borders.find(qn(f"w:{border_name}"))
        if element is None:
            element = OxmlElement(f"w:{border_name}")
            tbl_borders.append(element)
        if border_name in {"top", "bottom", "insideH"}:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "8")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")
        else:
            element.set(qn("w:val"), "nil")


def load_rows():
    data = json.loads(JSON_PATH.read_text(encoding="utf-8"))
    rows = []
    for endpoint, items in data.items():
        for item in items:
            rows.append(
                [
                    endpoint,
                    str(item["concurrency"]),
                    str(item["requests"]),
                    f'{item["avg_ms"]:.2f}',
                    f'{item["p95_ms"]:.2f}',
                    f'{item["error_rate"] * 100:.2f}%',
                ]
            )
    return rows


def find_endpoint_rows(data, endpoint_name: str):
    return data.get(endpoint_name, [])


def endpoint_summary(rows):
    if not rows:
        return None
    return {
        "min_avg": min(item["avg_ms"] for item in rows),
        "max_avg": max(item["avg_ms"] for item in rows),
        "min_p95": min(item["p95_ms"] for item in rows),
        "max_p95": max(item["p95_ms"] for item in rows),
        "max_error_rate": max(item["error_rate"] for item in rows),
    }


def main():
    data = json.loads(JSON_PATH.read_text(encoding="utf-8"))
    rows = load_rows()
    product_rows = find_endpoint_rows(data, "商品查询接口")
    order_rows = find_endpoint_rows(data, "订单提交接口")
    recommendation_rows = find_endpoint_rows(data, "推荐接口")

    product_summary = endpoint_summary(product_rows)
    order_summary = endpoint_summary(order_rows)
    recommendation_summary = endpoint_summary(recommendation_rows)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)
    section.start_type = WD_SECTION.CONTINUOUS

    add_text_paragraph(
        doc,
        "4.3 性能测试结果（可直接插入正文）",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        east_asia="黑体",
        size=14,
        before=0,
        after=12,
    )

    add_text_paragraph(
        doc,
        "为进一步验证系统在中低并发场景下的运行稳定性，本文选取商品查询接口、订单提交接口和推荐接口作为核心测试对象，并在并发用户数分别为10、30和50的条件下开展接口性能测试。测试过程中分别记录请求次数、平均响应时间、95%响应时间和错误率等指标，以分析系统在不同负载下的响应能力与稳定性。如表4-x所示，本次测试共覆盖三类核心接口。",
        size=10,
        before=0,
        after=6,
    )

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(caption, before=12, after=0, line=20)
    run = caption.add_run("表4-x 系统核心接口性能测试结果")
    run.bold = True
    set_font(run, east_asia="黑体", western="Times New Roman", size=10)

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["接口名称", "并发用户数", "请求次数", "平均响应时间/ms", "95%响应时间/ms", "错误率"]
    for i, text in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], text, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)

    set_three_line_table(table)

    add_text_paragraph(
        doc,
        (
            "如表4-x所示，在并发用户数分别为10、30和50的条件下，"
            f"商品查询接口平均响应时间由{product_summary['min_avg']:.2f} ms上升至{product_summary['max_avg']:.2f} ms，"
            f"95%响应时间由{product_summary['min_p95']:.2f} ms上升至{product_summary['max_p95']:.2f} ms，"
            f"错误率最高为{product_summary['max_error_rate'] * 100:.2f}%；"
            f"推荐接口平均响应时间由{recommendation_summary['min_avg']:.2f} ms上升至{recommendation_summary['max_avg']:.2f} ms，"
            f"95%响应时间由{recommendation_summary['min_p95']:.2f} ms上升至{recommendation_summary['max_p95']:.2f} ms，"
            f"错误率最高为{recommendation_summary['max_error_rate'] * 100:.2f}%。"
            "上述结果说明商品查询接口与推荐接口在当前测试规模下均能够稳定返回结果，且随着并发增加，推荐接口的耗时增长相对更明显。"
        ),
        size=10,
        before=12,
        after=6,
    )

    order_paragraph = (
        "订单提交接口在本轮测试中表现稳定。"
        f"在并发用户数分别为10、30和50的条件下，其平均响应时间介于{order_summary['min_avg']:.2f} ms至{order_summary['max_avg']:.2f} ms之间，"
        f"95%响应时间介于{order_summary['min_p95']:.2f} ms至{order_summary['max_p95']:.2f} ms之间，"
        f"错误率最高为{order_summary['max_error_rate'] * 100:.2f}%。"
        "说明在优化订单号生成策略后，订单写入链路已能够较好支撑当前中低并发场景下的提交请求。"
    )
    if order_summary["max_error_rate"] > 0:
        order_paragraph = (
            "订单提交接口在本轮测试中整体可用，但仍存在少量失败请求。"
            f"其平均响应时间介于{order_summary['min_avg']:.2f} ms至{order_summary['max_avg']:.2f} ms之间，"
            f"95%响应时间介于{order_summary['min_p95']:.2f} ms至{order_summary['max_p95']:.2f} ms之间，"
            f"错误率最高为{order_summary['max_error_rate'] * 100:.2f}%。"
            "相比优化前由订单号重复引发的大规模失败，当前结果表明订单唯一性问题已经得到明显改善，但后续仍可继续从数据库写入效率和事务链路稳定性方面进一步优化。"
        )

    add_text_paragraph(
        doc,
        order_paragraph,
        size=10,
        before=0,
        after=0,
    )

    doc.save(OUTPUT_PATH)
    print(str(OUTPUT_PATH))


if __name__ == "__main__":
    main()
