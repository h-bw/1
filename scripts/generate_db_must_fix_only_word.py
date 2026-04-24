from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


OUTPUT = r"D:\DevelopSoft\lingshi\数据库相关必改内容.docx"


def set_run_font(run, size=12, east_asia="宋体", western="Times New Roman", bold=False):
    run.bold = bold
    run.font.name = western
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=20):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = Pt(line)


def add_para(doc, text, *, bold=False, size=12, east_asia="宋体", align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0):
    p = doc.add_paragraph()
    set_paragraph(p, align=align, before=before, after=after)
    r = p.add_run(text)
    set_run_font(r, size=size, east_asia=east_asia, bold=bold)
    return p


def main():
    doc = Document()

    add_para(
        doc,
        "数据库相关必改内容",
        bold=True,
        size=16,
        east_asia="黑体",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=6,
    )
    add_para(
        doc,
        "以下内容仅包含建议直接替换的数据库相关正文，重点用于统一 3.2.2、3.2.3、3.2.4 与 3.3 的口径。",
        size=12,
    )

    add_para(doc, "一、3.2.2 第一段替换稿", bold=True, size=14, east_asia="黑体", align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=3)
    add_para(
        doc,
        "数据库逻辑结构设计是在数据库概念结构设计的基础上，将 E-R 图中的实体、属性以及实体之间的联系转换为关系数据库中的关系模式。关系模式是关系数据库逻辑结构的基本表示形式，主要由关系名、属性集合、主码以及外码约束组成。本系统根据零食商城小程序的业务需求，将用户、收货地址、商品分类、商品信息等实体，以及由实体联系转换得到的购物车关系和订单购买关系，进一步整理为相应的关系模式，并通过主码和外码建立各关系之间的联系。",
        size=12,
    )

    add_para(doc, "二、3.2.3 数据库关系图说明替换稿", bold=True, size=14, east_asia="黑体", align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=3)
    add_para(
        doc,
        "核心交易模块数据库关系图主要描述系统中与交易流程密切相关的数据对象及其联系，包括用户表、收货地址表、商品分类表、商品信息表、购物车表以及订单购买关系等。其中，用户与收货地址之间是一对多关系，商品分类与商品信息之间是一对多关系，用户与商品之间则通过购物车关系和订单购买关系建立联系。该图能够较为清晰地反映系统在概念层和逻辑层的数据组织方式，为后续数据库表设计与业务实现提供依据。",
        size=12,
    )
    add_para(
        doc,
        "需要说明的是，在系统实际实现中，订单购买关系并不是通过单一物理表直接保存，而是由订单主表、订单明细表以及商品快照信息共同支撑，并通过兼容视图进行统一表示。因此，在数据库关系图说明中，若从逻辑抽象层面进行描述，应统一使用“订单购买关系”这一表述；若从系统实现层面展开说明，则可进一步指出该关系由订单主表与订单明细表组合实现。",
        size=12,
    )
    add_para(
        doc,
        "推荐模块数据库关系图则主要反映推荐计算所依赖的数据对象及其联系，包括用户表、订单购买关系、商品信息表和商品分类表等。推荐模块通过读取用户历史购买记录，结合商品信息构造用户—商品交互关系，并在此基础上计算用户偏好和候选推荐结果。因此，该图在逻辑层面上应统一描述为“订单购买关系参与推荐计算”，而在实现层面上则对应订单主表与订单明细表组合形成的订单数据来源。",
        size=12,
    )

    add_para(doc, "三、3.2.4 订单结构说明替换稿", bold=True, size=14, east_asia="黑体", align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=3)
    add_para(
        doc,
        "订单购买表用于表示用户与商品之间在下单场景下形成的购买关系，其结构如表3-7所示。为与前文概念结构设计和逻辑结构设计保持一致，本文在数据库表设计部分统一采用“订单购买表”这一逻辑表达。在系统实际实现中，该关系并非由单一物理表直接存储，而是由订单主表、订单明细表及商品快照信息共同支撑，并通过 order_purchase 兼容视图进行统一表示。这样既能保持论文中 E-R 图、关系模式和数据库表设计之间的一致性，又不会破坏现有系统的订单处理流程。",
        size=12,
    )
    add_para(
        doc,
        "其中，order_id 用于标识订单记录，user_id 用于关联下单用户，product_id 用于标识订单中的商品，create_time 表示下单时间，status 表示订单状态，total_amount 表示订单总金额，quantity 和 price 分别表示购买数量与成交价格。该表主要体现用户与商品之间的订单购买联系，也为订单查询、推荐分析以及后续统计处理提供基础数据支持。",
        size=12,
    )

    add_para(doc, "四、3.3 订单处理模块设计替换稿", bold=True, size=14, east_asia="黑体", align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=3)
    add_para(
        doc,
        "订单处理模块是商城交易闭环的重要组成部分，用于将购物车中待结算的商品转化为可查询的订单信息。在逻辑抽象层面，该模块对应用户与商品之间形成的订单购买关系；在系统实现层面，则采用订单主表与订单明细表相结合的主从结构完成订单落库。用户在购物车中选择待购买商品后进入订单确认页面，系统会对收货地址、商品数量和订单总价进行检查，检查无误后由前端向后台提交订单请求，在事务控制下完成订单主信息保存、订单商品明细写入以及购物车数据清理，最终生成可查询的订单结果并返回前端页面。",
        size=12,
    )
    add_para(
        doc,
        "从具体实现来看，订单主表主要保存订单编号、用户编号、收货信息、订单总金额、商品总件数、订单状态和下单时间等基础信息；订单明细表则保存订单中各商品的名称、价格、数量、图片及商品编号等快照数据。为了保证交易链路的一致性，后端将“主表写入 + 明细写入”放在同一事务中进行处理，避免出现仅保存部分订单信息的情况。同时，系统通过优化订单号生成策略提升了订单编号在并发场景下的唯一性和稳定性，从而保证订单处理流程能够正常运行。",
        size=12,
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
